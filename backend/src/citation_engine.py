# -*- coding: utf-8 -*-
"""
Citation Engine - Moteur de Citations Academiques
Moteur de Rédaction Magique v5.0
Rev. Alphonse Owoudou, PhD

Formats: APA 7th, Chicago/Turabian, MLA, BibTeX
Gere les sources classiques ET les contenus generes par IA.
"""

from datetime import datetime
from typing import List, Dict, Optional
import json
import logging
import io

logger = logging.getLogger(__name__)

THINKERS_DB = {
    "don_bosco": {
        "full_name": "Giovanni Bosco",
        "birth_year": 1815, "death_year": 1888,
        "key_works": [
            {"title": "Il sistema preventivo nell'educazione della gioventù", "year": 1877, "type": "treatise"},
            {"title": "Memorie dell'Oratorio di S. Francesco di Sales", "year": 1875, "type": "autobiography"},
            {"title": "Lettera da Roma", "year": 1884, "type": "letter"},
        ]
    },
    "maria_montessori": {
        "full_name": "Maria Montessori",
        "birth_year": 1870, "death_year": 1952,
        "key_works": [
            {"title": "Il metodo della pedagogia scientifica", "year": 1909, "type": "book"},
            {"title": "La scoperta del bambino", "year": 1950, "type": "book"},
            {"title": "La mente del bambino", "year": 1949, "type": "book"},
        ]
    },
    "paulo_freire": {
        "full_name": "Paulo Freire",
        "birth_year": 1921, "death_year": 1997,
        "key_works": [
            {"title": "Pedagogia do Oprimido", "year": 1968, "type": "book"},
            {"title": "Educação como Prática da Liberdade", "year": 1967, "type": "book"},
            {"title": "Pedagogia da Autonomia", "year": 1996, "type": "book"},
        ]
    },
    "hannah_arendt": {
        "full_name": "Hannah Arendt",
        "birth_year": 1906, "death_year": 1975,
        "key_works": [
            {"title": "The Human Condition", "year": 1958, "type": "book"},
            {"title": "The Crisis in Education", "year": 1961, "type": "essay"},
            {"title": "Between Past and Future", "year": 1961, "type": "book"},
        ]
    },
    "tommaso_aquino": {
        "full_name": "Tommaso d'Aquino",
        "birth_year": 1225, "death_year": 1274,
        "key_works": [
            {"title": "Summa Theologiae", "year": 1274, "type": "treatise"},
            {"title": "De Magistro", "year": 1256, "type": "treatise"},
            {"title": "Quaestiones Disputatae De Veritate", "year": 1259, "type": "treatise"},
        ]
    },
    "jean_piaget": {
        "full_name": "Jean Piaget",
        "birth_year": 1896, "death_year": 1980,
        "key_works": [
            {"title": "La construction du réel chez l'enfant", "year": 1937, "type": "book"},
            {"title": "La psychologie de l'intelligence", "year": 1947, "type": "book"},
            {"title": "L'épistémologie génétique", "year": 1970, "type": "book"},
        ]
    },
    "lev_vygotsky": {
        "full_name": "Lev Vygotsky",
        "birth_year": 1896, "death_year": 1934,
        "key_works": [
            {"title": "Pensiero e linguaggio", "year": 1934, "type": "book"},
            {"title": "Mind in Society", "year": 1978, "type": "book"},
        ]
    },
    "edith_stein": {
        "full_name": "Edith Stein",
        "birth_year": 1891, "death_year": 1942,
        "key_works": [
            {"title": "Il problema dell'empatia", "year": 1917, "type": "dissertation"},
            {"title": "Scientia Crucis", "year": 1942, "type": "book"},
        ]
    },
}


def format_citation(source: Dict, style: str = "apa") -> str:
    formatters = {
        "apa": _format_apa,
        "chicago": _format_chicago,
        "mla": _format_mla,
        "bibtex": _format_bibtex,
    }
    formatter = formatters.get(style, _format_apa)
    return formatter(source)


def _format_apa(source: Dict) -> str:
    authors = source.get("authors", [])
    year = source.get("year", "n.d.")
    title = source.get("title", "Untitled")
    source_type = source.get("type", "book")
    url = source.get("url", "")
    publisher = source.get("publisher", "")

    if authors:
        author_str = ", ".join(
            f"{a.split()[-1]}, {a.split()[0][0]}." if " " in a else a
            for a in authors[:3]
        )
        if len(authors) > 3:
            author_str += " et al."
    else:
        author_str = "Autore sconosciuto"

    citation = f"{author_str} ({year}). *{title}*."
    if publisher:
        citation += f" {publisher}."
    if url:
        citation += f" {url}"
    return citation


def _format_chicago(source: Dict) -> str:
    authors = source.get("authors", [])
    year = source.get("year", "n.d.")
    title = source.get("title", "Untitled")
    publisher = source.get("publisher", "")

    if authors:
        first = authors[0]
        parts = first.split()
        author_str = f"{parts[-1]}, {' '.join(parts[:-1])}" if len(parts) > 1 else first
        for a in authors[1:3]:
            author_str += f", and {a}"
    else:
        author_str = "Autore sconosciuto"

    citation = f"{author_str}. *{title}*."
    if publisher:
        citation += f" {publisher},"
    citation += f" {year}."
    return citation


def _format_mla(source: Dict) -> str:
    authors = source.get("authors", [])
    title = source.get("title", "Untitled")
    publisher = source.get("publisher", "")
    year = source.get("year", "n.d.")

    if authors:
        first = authors[0]
        parts = first.split()
        author_str = f"{parts[-1]}, {' '.join(parts[:-1])}" if len(parts) > 1 else first
    else:
        author_str = "Autore sconosciuto"

    citation = f'{author_str}. *{title}*.'
    if publisher:
        citation += f" {publisher},"
    citation += f" {year}."
    return citation


def _format_bibtex(source: Dict) -> str:
    authors = source.get("authors", [])
    year = source.get("year", "0000")
    title = source.get("title", "Untitled")
    key = title.split()[0].lower() + str(year) if title else "ref"

    return (
        f"@book{{{key},\n"
        f"  author = {{{' and '.join(authors) if authors else 'Unknown'}}},\n"
        f"  title = {{{title}}},\n"
        f"  year = {{{year}}}\n"
        f"}}"
    )


def generate_bibliography(sources: List[Dict], style: str = "apa") -> str:
    citations = [format_citation(s, style) for s in sources]
    citations.sort()
    header = {
        "apa": "Riferimenti bibliografici",
        "chicago": "Bibliografia",
        "mla": "Works Cited",
        "bibtex": "% BibTeX Bibliography",
    }.get(style, "Bibliografia")

    if style == "bibtex":
        return "\n\n".join(citations)

    numbered = [f"{i+1}. {c}" for i, c in enumerate(citations)]
    return f"**{header}**\n\n" + "\n".join(numbered)


def create_ai_provenance_note(
    model: str,
    feature: str,
    sources_used: List[str],
    date: Optional[str] = None,
) -> str:
    if not date:
        date = datetime.now().strftime("%d %B %Y")
    sources_str = "; ".join(sources_used) if sources_used else "corpus interne"
    return (
        f"[Nota di provenienza IA] Contenuto generato da Moteur Magique v5.0 "
        f"({model}, {date}) tramite la funzione '{feature}', "
        f"basato su: {sources_str}. "
        f"Il contenuto è stato prodotto da intelligenza artificiale e deve essere "
        f"verificato criticamente prima dell'uso in pubblicazioni accademiche."
    )


def create_export_document(
    title: str,
    content: str,
    bibliography: str,
    provenance_note: str,
    format_type: str = "docx",
) -> Optional[bytes]:
    if format_type == "docx":
        return _export_docx(title, content, bibliography, provenance_note)
    return None


def _export_docx(title: str, content: str, bibliography: str, provenance_note: str) -> Optional[bytes]:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")

        for para_text in content.split("\n\n"):
            if para_text.strip():
                if para_text.strip().startswith("##"):
                    doc.add_heading(para_text.strip().lstrip("#").strip(), level=2)
                elif para_text.strip().startswith("**") and para_text.strip().endswith("**"):
                    p = doc.add_paragraph()
                    run = p.add_run(para_text.strip().strip("*"))
                    run.bold = True
                else:
                    doc.add_paragraph(para_text.strip())

        doc.add_page_break()
        doc.add_heading("Bibliografia", level=2)
        for line in bibliography.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip(), style='List Number')

        doc.add_paragraph("")
        note_para = doc.add_paragraph()
        note_run = note_para.add_run(provenance_note)
        note_run.italic = True
        note_run.font.size = Pt(10)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except ImportError:
        logger.warning("python-docx not installed, DOCX export unavailable")
        return None
    except Exception as e:
        logger.error(f"Error creating DOCX: {e}")
        return None


def get_thinkers_list() -> List[Dict]:
    return [
        {
            "id": tid,
            "name": t["full_name"],
            "period": f"{t['birth_year']}-{t['death_year']}",
            "works_count": len(t["key_works"]),
        }
        for tid, t in THINKERS_DB.items()
    ]


def get_thinker_info(thinker_id: str) -> Optional[Dict]:
    return THINKERS_DB.get(thinker_id)


def add_custom_thinker(
    thinker_id: str,
    full_name: str,
    birth_year: int,
    death_year: int,
    description: str,
    key_works: List[Dict],
) -> Dict:
    THINKERS_DB[thinker_id] = {
        "full_name": full_name,
        "birth_year": birth_year,
        "death_year": death_year,
        "key_works": key_works,
        "custom": True,
    }
    _save_custom_thinkers()
    return {"success": True, "id": thinker_id, "name": full_name}


def delete_custom_thinker(thinker_id: str) -> Dict:
    info = THINKERS_DB.get(thinker_id)
    if not info:
        return {"success": False, "error": "Pensatore non trovato"}
    if not info.get("custom"):
        return {"success": False, "error": "Impossibile eliminare un pensatore predefinito"}
    del THINKERS_DB[thinker_id]
    _save_custom_thinkers()
    return {"success": True}


def _get_custom_path():
    import os
    data_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data")
    os.makedirs(data_dir, exist_ok=True)
    return os.path.join(data_dir, "custom_thinkers.json")


def _save_custom_thinkers():
    custom = {k: v for k, v in THINKERS_DB.items() if v.get("custom")}
    try:
        with open(_get_custom_path(), "w", encoding="utf-8") as f:
            json.dump(custom, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Error saving custom thinkers: {e}")


def _load_custom_thinkers():
    try:
        path = _get_custom_path()
        import os
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                custom = json.load(f)
                THINKERS_DB.update(custom)
                logger.info(f"Loaded {len(custom)} custom thinkers")
    except Exception as e:
        logger.error(f"Error loading custom thinkers: {e}")


_load_custom_thinkers()
